import datetime
from docx import Document

def generate_week_text(week_number, year):
    doc = Document()

    # Add the title
    title = doc.add_heading(f'WEEK {week_number}', level=1)

    # Add the week goals section
    doc.add_heading('Week goals', level=2)

    # Calculate the start date of the week
    start_date = datetime.date.fromisocalendar(year, week_number, 1)

    for i in range(7):
        current_date = start_date + datetime.timedelta(days=i)
        formatted_date = current_date.strftime("%a %d.%m.%y")
        doc.add_heading(formatted_date, level=2)

    # Add the week recap section
    doc.add_heading('Week recap', level=2)

    return doc

def main():
    user_input = input("Enter week number(s) (e.g., '38', '38-40', or '38-'): ")

    # Split the input into parts based on '-'
    input_parts = user_input.split('-')

    if len(input_parts) == 1:
        # Single week input
        week_numbers = [int(input_parts[0])]
    elif len(input_parts) == 2 and input_parts[1].strip() == "":
        # Input with a trailing dash, generate all weeks from the specified week to the end of the year
        start_week = int(input_parts[0])
        current_year = datetime.datetime.now().year
        week_number_of_dec_31 = datetime.datetime(current_year, 12, 31).isocalendar()[1]
        week_number_of_week_before_dec_31_week = datetime.datetime(current_year, 12, 31 - 7).isocalendar()[1]
        week_numbers = list(range(start_week, (week_number_of_week_before_dec_31_week if week_number_of_dec_31 == 1 else week_number_of_dec_31) + 1))
    elif len(input_parts) == 2:
        # Range of weeks input
        start_week = int(input_parts[0])
        end_week = int(input_parts[1])
        week_numbers = list(range(start_week, end_week + 1))
    else:
        print("Invalid input format. Please enter a single week, a range of weeks, or a range with a trailing dash.")
        return

    year = int(input("Enter the year: "))

    # Create a single document and add weeks to it
    doc = Document()
    for week_number in week_numbers:
        week_doc = generate_week_text(week_number, year)
        for element in week_doc.element.body:
            doc.element.body.append(element)

    # Save the single document, overwriting the previous content
    doc.save('Weeks.docx')

if __name__ == "__main__":
    main()
